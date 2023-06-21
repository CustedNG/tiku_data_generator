import os
import sys
import docx


def convert_dir(doc_dir):
    '''
    param doc_dir: 文档目录, 如'xxx/'
    return: None
    '''
    skip_files = []
    for i in os.listdir(doc_dir):
        if not i.endswith('.docx'):
            if not i.endswith('.txt') and not i.endswith('.json') and not i == '.DS_Store':
                skip_files.append(i)
            continue
        if not doc_dir.endswith('/'):
            doc_dir = doc_dir + '/'
        f = open(doc_dir + i.replace('.docx', '.txt'), 'w')
        # 打开docx的文档并读入名为file的变量中
        file = docx.Document(doc_dir + i)
        # 输入docx中的段落数，以检查是否空文档
        para_len = len(file.paragraphs)
        if para_len < 10:
            print('段落数小于10，跳过')
            continue
        # 将每个段落的内容都写进去txt里面
        for para in file.paragraphs:
            f.write('\n')
            f.write(para.text)
        f.close()

    if len(skip_files) > 0:
        raise Exception('\n以下文件由于格式问题，没有转换：\n' + str(skip_files))


if __name__ == '__main__':
    args = sys.argv
    if len(args) < 2:
        print('请输入文件夹路径')
        exit(1)
    convert_dir(args[1])