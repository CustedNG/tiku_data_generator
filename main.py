from typing import List

import json
import re
import os
import docx

answer_convert_map = {
    'A': 0,
    'B': 1,
    'C': 2,
    'D': 3,
    'E': 4,
    'F': 5,
    'G': 6,
    'H': 7,
    '是': 0,
    '对': 0,
    '正确': 0,
    '错': 1,
    '否': 1,
    '错误': 1,
}


def parse_not_implement():
    raise NotImplementedError()


# [doc_dir]: 'xxx/'
def docx2txt(doc_dir):
    skip_files = []
    for i in os.listdir(doc_dir):
        if not i.endswith('.docx'):
            if not i.endswith('.txt') and not i.endswith('.json') and not i == '.DS_Store':
                skip_files.append(i)
            continue
        f = open(doc_dir + i.replace('.docx', '.txt'), 'w')
        # 打开docx的文档并读入名为file的变量中
        file = docx.Document(doc_dir + i)
        # 输入docx中的段落数，以检查是否空文档
        para_len = len(file.paragraphs)
        if para_len < 10:
            print('段落数小于10，跳过')
            continue
        # 将每个段落的内容都写进去txt里面
        for para in file.paragraphs:
            f.write('\n')
            f.write(para.text)
        f.close()
        print('转换完成：' + i)

    if len(skip_files) > 0:
        raise Exception('\n以下文件由于格式问题，没有转换：\n' + str(skip_files))
    return


def parse_maogai():
    doc_dir = 'maogai/'
    difficulty_locator = '难易'
    answer_locator = '答案'

    # docx2txt(doc_dir)
    for file in os.listdir(doc_dir):
        if not file.endswith('.txt'):
            continue
        # 题目列表，纯str
        raw_list: List[str] = []
        idx = 0
        f = open(doc_dir + '/' + file, 'r')
        lines = f.readlines()
        f.close()
        for i in range(len(lines)):
            lines[i] = lines[i].replace(' ', ' ')
            # 根据题号分割题目
            if re.match(r'( )*[1-9]+.*', lines[i]):
                raw_list.append(lines[idx:i])
                idx = i
        raw_list.append(lines[idx:])
        raw_list = raw_list[1:]

        # 标准结构的题目列表
        ti_list = []
        for ti_raw in raw_list:
            # 解析标题
            ti_head = re.match(r'( )*[1-9]+[0-9]*( )*(、|\.|．)*( )*', ti_raw[0])
            title = ti_raw[0].replace(ti_head.group(), '').strip()

            # 解析答案
            answer_raw = ''
            for item in ti_raw:
                if answer_locator in item:
                    answer_raw = item.strip()
            answer = []
            for separator in ['】', '：']:
                if separator in answer_raw:
                    answer_raw = answer_raw.split(separator)[1]
            if answer_raw in answer_convert_map:
                answer.append(answer_convert_map[answer_raw])
            else:
                for item in answer_raw:
                    if item not in answer_convert_map:
                        continue
                    answer.append(answer_convert_map[item])

            # 解析选项
            options = []
            answer_idx = -2
            idx_num = 0
            for li in ti_raw:
                if answer_locator in li:
                    answer_idx = idx_num
                idx_num += 1
            for option in ti_raw[1:answer_idx]:
                if answer_locator in option or difficulty_locator in option or not option:
                    continue
                temp_ops = option.strip().split(' ')
                for item in temp_ops:
                    if re.match(r'(A|B|C|D|E|F|G)', item):
                        temp = item.strip()[1:]
                        for it in ['、', '.', '．']:
                            if temp.startswith(it):
                                temp = temp.replace(it, '')
                        options.append(temp.strip())

            # 解析类型
            ti_type = 0
            answer_len = len(answer)
            if answer_len > 1:
                ti_type = 1
            if len(options) == 0 and answer_len == 1:
                options = ['对', '错']

            # 添加解析完成的题目到列表
            ti_list.append({
                'title': title,
                'options': options,
                'answer': answer,
                'type': ti_type
            })

        print(f'{file}: {len(ti_list)}题')
        print(f'{ti_list[0]}...')

        with open(doc_dir + file.replace('.txt', '.json'), 'w') as f:
            f.write(json.dumps(ti_list, ensure_ascii=False, indent=4))


subject_map = {
    'maogai': {'name': '毛概', 'func': parse_maogai},
    'clang1': {'name': 'C语言上', 'func': parse_not_implement},
    'clang2': {'name': 'C语言下', 'func': parse_not_implement},
    'junli1': {'name': '军理上', 'func': parse_not_implement},
    'junli2': {'name': '军理下', 'func': parse_not_implement},
    'jinxiandaishi': {'name': '近代史', 'func': parse_not_implement},
    'sixiu': {'name': '思修', 'func': parse_not_implement},
    'makesi': {'name': '马克思', 'func': parse_not_implement},
}

if __name__ == '__main__':
    subjects = subject_map.keys()
    for subject in subjects:
        if os.path.exists(subject):
            info = subject_map[subject]
            print(f'开始解析{info["name"]}')
            info['func']()
